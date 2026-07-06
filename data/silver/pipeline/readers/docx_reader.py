import hashlib
import io
import os
from datetime import datetime, timezone

from data.silver.pipeline.readers.base import BaseReader
from data.silver.schemas.source import SourceData


class DocxReader(BaseReader):
    def can_handle(self, source: SourceData) -> bool:
        ext = os.path.splitext(source.filename.lower())[1]
        return ext in (".docx", ".doc")

    def read(self, source: SourceData) -> dict:
        try:
            from docx import Document
            doc = Document(io.BytesIO(source.raw_data))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            tables = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    tables.append(" | ".join(cells))
            content = "\n".join(paragraphs)
            if tables:
                content += "\n\n" + "\n".join(tables)
        except Exception:
            content = ""
        return {"content": content}

    def extract_metadata(self, source: SourceData) -> dict:
        ext = os.path.splitext(source.filename.lower())[1]
        checksum = hashlib.sha256(source.raw_data).hexdigest()
        created_time = None
        modified_time = None
        try:
            from docx import Document
            doc = Document(io.BytesIO(source.raw_data))
            props = doc.core_properties
            created_time = props.created if hasattr(props, "created") else None
            modified_time = props.modified if hasattr(props, "modified") else None
        except Exception:
            pass
        return {
            "extension": ext,
            "checksum": checksum,
            "size_bytes": source.size_bytes,
            "object_key": source.object_key,
            "bucket": source.bucket,
            "created_time": created_time,
            "modified_time": modified_time,
            "source": f"minio://{source.bucket}/{source.object_key}",
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
